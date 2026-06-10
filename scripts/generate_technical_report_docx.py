"""將技術報告內容輸出為 Word 章節結構（.docx）。"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "技術報告.docx"


def set_run_font(run, size_pt: int = 12, bold: bool = False, color: RGBColor | None = None):
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_toc_field(doc: Document) -> None:
    """插入 Word 目錄欄位（開啟後按 F9 或「更新目錄」）。"""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    run2 = p.add_run("（請在 Word 中按 F9 或右鍵→更新欄位以產生目錄）")
    set_run_font(run2, 10, color=RGBColor(0x66, 0x66, 0x66))

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run2._r.append(fld_end)


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, {1: 16, 2: 14, 3: 12}.get(level, 12), bold=True)


def add_para(doc: Document, text: str, bold: bool = False, indent: bool = False) -> None:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, bold=bold)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, 10)
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(line)
        set_run_font(run, 10)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def build_document() -> Document:
    doc = Document()

    # 頁面邊界
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.5)

    # ===== 封面 =====
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("輕量智敘（Light Narrative）\n技術報告")
    set_run_font(tr, 22, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("本地優先之中文敘事一致性檢查系統 — 期末 MVP 技術文件")
    set_run_font(sr, 14)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        "版本：期末 MVP（2026-06 交付版）\n"
        "文件性質：系統設計、實作說明與驗收紀錄\n"
        "最後驗證：python -m pytest tests/ -q → 30 passed"
    )
    set_run_font(mr, 11)

    doc.add_page_break()

    # ===== 目錄 =====
    add_heading(doc, "目錄", 1)
    add_para(doc, "以下目錄欄位將於 Word 開啟後自動產生（若未顯示，請更新欄位）：")
    add_toc_field(doc)
    doc.add_page_break()

    # ===== 第 1 章 =====
    add_heading(doc, "第 1 章　摘要", 1)
    add_para(
        doc,
        "輕量智敘是一套本地優先的中文敘事一致性檢查工具，協助作者在小說、劇本與原創 IP 文本中"
        "偵測敘事矛盾——例如角色死亡後仍活動、唯一物品矛盾、世界規則違反、人設漂移與設定反轉。",
        indent=True,
    )
    add_para(doc, "本系統不是自動寫作或自動修稿工具。其核心價值在於：")
    add_bullets(
        doc,
        [
            "以 7 條確定性（deterministic）衝突規則產出可重現、可測試的檢查結果",
            "以 Heuristic 抽取為主，可選 Ollama LLM 輔助實體／事件抽取",
            "以 Streamlit 單頁 UI 提供上傳、分析、知識圖與 JSON 報告匯出",
            "以 SQLite 在本機持久化專案資料",
        ],
    )
    add_para(
        doc,
        "衝突判斷不由 LLM 決定，確保結果穩定、適合期末展示與單元測試驗收。",
        indent=True,
    )

    # ===== 第 2 章 =====
    add_heading(doc, "第 2 章　專案背景與目標", 1)
    add_heading(doc, "2.1 問題陳述", 2)
    add_para(doc, "長篇敘事創作中，作者常面臨：")
    add_bullets(
        doc,
        [
            "角色狀態前後不一致（死亡後仍說話、行動）",
            "世界規則與後續事件矛盾（「亡者不能復活」卻出現復活）",
            "唯一物品設定被違反（「僅有一把鑰匙」卻出現備用品）",
            "人設與行為不符（膽小角色突然展現高能力）",
            "早期世界設定與後期敘述反轉卻無伏筆鋪陳",
        ],
    )
    add_para(
        doc,
        "人工逐章比對耗時且易漏。本專題目標是建立一個可本地執行、規則可解釋的 MVP，"
        "輔助作者發現問題並提供 suggested_fix 文字建議。",
        indent=True,
    )

    add_heading(doc, "2.2 設計原則", 2)
    add_table(
        doc,
        ["原則", "說明"],
        [
            ["本地優先", "資料存於本機 SQLite，無需雲端帳號"],
            ["規則驅動", "衝突偵測由 conflict_rules.py 決定，非黑箱 LLM"],
            ["LLM 可選", "Ollama 僅用於抽取層；未連線自動 fallback"],
            ["可測試", "每條規則皆有對應 pytest；另有 evaluation 模組"],
            ["誠實邊界", "未實作功能明確列於 LIMITATIONS.md"],
        ],
    )

    add_heading(doc, "2.3 非目標（刻意不在本期範圍）", 2)
    add_bullets(
        doc,
        [
            "FastAPI 後端、登入、雲端協作",
            "HuggingFace 本地推理（目前為 stub）",
            "獨立時間線引擎與視圖",
            "自動修稿（僅輸出建議文字）",
        ],
    )

    # ===== 第 3 章 =====
    add_heading(doc, "第 3 章　系統架構", 1)
    add_heading(doc, "3.1 整體架構", 2)
    add_para(doc, "系統採分層架構，資料流如下：", indent=True)
    add_code_block(
        doc,
        [
            "【展示層】app.py（Streamlit 單頁 UI）",
            "    ↓",
            "【分析管線】ingest → chunker → agents（抽取）→ state_tracker → conflict_rules → report_writer",
            "              ↘ graph_builder（知識圖）",
            "    ↓",
            "【持久化】SQLite（light_narrative.db）+ data/exports/（JSON 報告）",
            "【可選】Ollama LLM（抽取輔助）、PyVis（互動圖形）",
        ],
    )

    add_heading(doc, "3.2 分層職責", 2)
    add_table(
        doc,
        ["層級", "模組", "職責"],
        [
            ["展示層", "app.py", "專案管理、上傳、分析觸發、衝突報告、知識圖、JSON 匯出"],
            ["匯入層", "ingest.py", "解析 .txt / .md / .docx 為 Document"],
            ["分塊層", "chunker.py", "auto / chapter / scene / fixed 四種策略"],
            ["抽取層", "agents/", "Heuristic 或 Ollama 抽取實體與事件"],
            ["狀態層", "state_tracker.py", "角色死亡／活動、物品持有人追蹤"],
            ["規則層", "conflict_rules.py", "確定性衝突偵測與證據過濾"],
            ["輸出層", "report_writer.py, graph_builder.py", "結構化報告與 NetworkX 圖"],
            ["儲存層", "storage.py", "SQLite CRUD 與 JSON 匯出"],
            ["評估層", "evaluation/", "離線 ground truth 比對（不影響主系統）"],
        ],
    )

    add_heading(doc, "3.3 Agent 編排", 2)
    add_para(doc, "完整分析流程由 src/agents/report_agent.py 的 generate_full_report() 編排：")
    add_code_block(
        doc,
        [
            "chunks",
            "  → run_extraction()              # 逐 chunk 抽取 + merge",
            "  → derive_character_states() / derive_item_states()",
            "  → run_conflict_detection()      # detect_all_conflicts()",
            "  → build_report()",
            "  → save_extraction() / save_report()",
        ],
    )

    # ===== 第 4 章 =====
    add_heading(doc, "第 4 章　技術棧", 1)
    add_table(
        doc,
        ["類別", "技術", "版本需求", "用途"],
        [
            ["語言", "Python 3", "—", "全系統"],
            ["Web UI", "Streamlit", "≥ 1.32", "單頁應用"],
            ["資料驗證", "Pydantic", "≥ 2.6", "Schema 定義"],
            ["圖論", "NetworkX", "≥ 3.2", "知識圖建構"],
            ["視覺化", "PyVis", "≥ 0.3.2", "互動圖（可選 fallback）"],
            ["文件解析", "python-docx", "≥ 1.1", ".docx 上傳"],
            ["HTTP", "httpx", "≥ 0.27", "Ollama API 呼叫"],
            ["儲存", "SQLite", "內建", "本地持久化"],
            ["測試", "pytest", "≥ 8.0", "單元與整合測試"],
            ["LLM（可選）", "Ollama", "—", "llama3.2 等本地模型"],
        ],
    )

    # ===== 第 5 章 =====
    add_heading(doc, "第 5 章　核心模組說明", 1)

    add_heading(doc, "5.1 文本匯入（ingest.py）", 2)
    add_bullets(
        doc,
        [
            "支援 .txt、.md、.docx",
            "上傳後整份文本寫入 documents.raw_text",
            "Streamlit 預設上傳上限 200 MB",
        ],
    )

    add_heading(doc, "5.2 分塊（chunker.py）", 2)
    add_table(
        doc,
        ["策略", "行為"],
        [
            ["auto", "先依「第 X 章」切章；章內有場景符再切；超長段 fallback"],
            ["chapter", "僅依章節標題；無標題則整篇一塊"],
            ["scene", "依場景分隔符；無分隔符 fallback 為 fixed"],
            ["fixed", "依段落累積，每塊上限 DEFAULT_CHUNK_MAX_CHARS = 2500"],
        ],
    )

    add_heading(doc, "5.3 抽取層（agents/）", 2)
    add_para(doc, "Heuristic 模式（預設）：", bold=True)
    add_bullets(
        doc,
        [
            "entity_agent.py：regex／關鍵字抽取角色、地點、物品、世界規則",
            "event_agent.py：動作句型抽取事件（每 chunk 最多 40 筆）",
            "narrative_patterns.py：中文敘事句型、否定句、實體名稱驗證",
        ],
    )
    add_para(doc, "Ollama 模式（可選）：", bold=True)
    add_bullets(
        doc,
        [
            "每 chunk 送前 3000 字給 LLM（entity + event 各 1 次請求）",
            "單次 timeout 120 秒",
            "is_available() 失敗時自動 fallback，不中斷流程",
        ],
    )
    add_para(doc, "重要：LLM 只影響抽取品質，不參與衝突判斷。", bold=True)

    add_heading(doc, "5.4 狀態追蹤（state_tracker.py）", 2)
    add_bullets(
        doc,
        [
            "derive_character_states()：依死亡句型、觀察者句型、復活句型追蹤角色狀態",
            "derive_item_states()：追蹤物品持有人與位置變化",
            "含代詞回溯、地名誤判過濾、證據有效性檢查",
        ],
    )

    add_heading(doc, "5.5 衝突規則（conflict_rules.py）", 2)
    add_para(
        doc,
        "detect_all_conflicts() 依序執行 7 條規則（含 1 條額外子規則），"
        "最後去重合併並依嚴重度排序：",
    )
    add_table(
        doc,
        ["規則", "函式", "衝突類型", "偵測內容"],
        [
            ["Rule 1", "rule1_dead_then_active", "character_state_conflict", "角色死亡／安葬後仍活動"],
            ["Rule 2", "rule2_dead_cannot_resurrect", "world_rule_violation", "亡者不能復活 vs 復活事件"],
            ["Rule 3", "rule3_unique_item", "unique_item_conflict", "唯一物品規則 vs 備用品"],
            ["Rule 4", "rule4_item_location_holder", "item_location_conflict", "物品持有人不一致"],
            ["Rule 5", "rule5_night_bell", "world_rule_violation", "夜晚敲鐘規則 vs 夜間敲鐘"],
            ["Rule 8", "rule8_voice_prohibition", "world_rule_violation", "發聲禁令 vs 違規發聲"],
            ["Rule 6", "rule6_character_drift", "character_consistency_drift", "人設限制 vs 高能力行為"],
            ["Rule 7", "rule7_world_setting", "world_setting_conflict", "早期設定 vs 後期反轉"],
        ],
    )

    add_heading(doc, "5.6 知識圖（graph_builder.py）", 2)
    add_bullets(
        doc,
        [
            "以 NetworkX MultiDiGraph 建構實體關係",
            "PyVis 產生互動 HTML；失敗時顯示統計與 fallback 提示",
            "節點類型：角色、地點、物品、事件、規則",
        ],
    )

    add_heading(doc, "5.7 儲存（storage.py）", 2)
    add_table(
        doc,
        ["資料表", "內容"],
        [
            ["projects", "專案名稱、描述、建立時間"],
            ["documents", "上傳檔名、原始文本"],
            ["chunks", "分塊文本與 metadata"],
            ["extractions", "抽取結果 JSON"],
            ["reports", "分析報告 JSON"],
        ],
    )
    add_para(doc, "JSON 匯出路徑：data/exports/report_{project_id前8字}.json")

    # ===== 第 6 章 =====
    add_heading(doc, "第 6 章　資料模型", 1)
    add_para(doc, "核心 schema 定義於 src/schemas.py（Pydantic v2）：")
    add_code_block(
        doc,
        [
            "Project → Document → Chunk[]",
            "                ↓",
            "         ExtractionResult",
            "           ├── characters[], locations[], objects[]",
            "           ├── events[], world_rules[]",
            "           ├── character_states[], item_states[]",
            "                ↓",
            "         ConflictReport[] → AnalysisReport",
        ],
    )
    add_para(
        doc,
        "ConflictReport 為對外展示與 JSON 匯出的核心結構；"
        "AnalysisReport 聚合統計（chunk 數、衝突數、抽取摘要）。",
        indent=True,
    )

    # ===== 第 7 章 =====
    add_heading(doc, "第 7 章　處理流程", 1)
    add_heading(doc, "7.1 使用者操作流程", 2)
    add_numbered(
        doc,
        [
            "側欄建立／選擇專案",
            "上傳 .txt / .md / .docx，選擇分塊策略",
            "預覽分塊 → 執行敘事分析",
            "檢視抽取統計、衝突報告、知識圖",
            "匯出 JSON 至 data/exports/",
        ],
    )

    add_heading(doc, "7.2 系統內部管線", 2)
    add_table(
        doc,
        ["步驟", "元件", "動作"],
        [
            ["1", "Streamlit", "接收上傳文本"],
            ["2", "Chunker", "chunk_document() → Chunk[]"],
            ["3", "Extraction Agents", "run_extraction() → ExtractionResult"],
            ["4", "State Tracker", "derive_*_states()"],
            ["5", "Conflict Rules", "detect_all_conflicts() → ConflictReport[]"],
            ["6", "Report Writer", "build_report() → AnalysisReport"],
            ["7", "SQLite", "save_extraction / save_report"],
            ["8", "Streamlit", "顯示報告與知識圖"],
        ],
    )

    # ===== 第 8 章 =====
    add_heading(doc, "第 8 章　效能與容量建議", 1)
    add_para(
        doc,
        "本 MVP 非為整本長篇小說設計。實測與展示主要使用 samples/ 短篇（約 300～400 字、3～5 chunk）。",
        indent=True,
    )
    add_table(
        doc,
        ["面向", "限制／建議"],
        [
            ["上傳", "Streamlit 預設 200 MB；整份讀入記憶體"],
            ["分塊", "無 chunk 數量上限；文本越長分析越慢（大致線性）"],
            ["Heuristic", "每 chunk 依序 regex；事件每 chunk ≤ 40 筆"],
            ["Ollama", "chunk 數 × 2 次 API；單次 120s timeout"],
            ["建議用量", "展示：300～400 字；開發測試：500 字～2 萬字"],
            ["不建議", "十萬字以上整本小說一次分析"],
        ],
    )

    # ===== 第 9 章 =====
    add_heading(doc, "第 9 章　測試與驗收", 1)
    add_heading(doc, "9.1 單元測試（pytest）", 2)
    add_para(doc, "目前 30 個測試全數通過，涵蓋：")
    add_table(
        doc,
        ["測試檔", "驗證重點"],
        [
            ["test_character_state_conflict.py", "Rule 1：死亡後活動"],
            ["test_rule2_resurrection_violation.py", "Rule 2：復活違規"],
            ["test_unique_item_conflict.py", "Rule 3：唯一物品"],
            ["test_rule4_item_holder_conflict.py", "Rule 4：持有人矛盾"],
            ["test_world_rule_violation.py", "Rule 5：夜間敲鐘"],
            ["test_character_drift.py", "Rule 6：人設漂移"],
            ["test_world_setting_conflict.py", "Rule 7：設定反轉"],
            ["test_no_false_positive_clean_story.py", "無矛盾文本零誤報"],
            ["test_renamed_story_generalization.py", "換名文本泛化"],
            ["test_chunker.py", "分塊策略"],
            ["test_narrative_patterns.py", "句型與實體過濾"],
            ["test_storage.py", "SQLite CRUD"],
            ["test_smoke_modules.py", "ingest / graph / report smoke"],
            ["test_evaluate_report.py", "評估模組"],
            ["test_round3_fixes.py", "迭代修正回歸"],
        ],
    )
    add_para(doc, "驗收指令：")
    add_code_block(
        doc,
        [
            "pip install -r requirements.txt",
            "python -m pytest tests/ -q",
            "streamlit run app.py",
        ],
    )

    add_heading(doc, "9.2 評估模組（evaluation/）", 2)
    add_para(
        doc,
        "離線比對系統 JSON 報告與人工 ground truth，"
        "計算 TP / FP / FN、Precision、Recall、F1、Duplicate Rate、Evidence Accuracy。",
        indent=True,
    )
    add_table(
        doc,
        ["文本", "預期", "輸出", "TP", "FP", "FN", "Precision", "Recall", "F1"],
        [
            ["fog_bell（霧城之鐘）", "6", "6", "6", "0", "0", "1.0", "1.0", "1.0"],
            ["clean_story（無矛盾）", "0", "0", "0", "0", "0", "—", "—", "—"],
            ["renamed_story（換名）", "5", "7", "5", "2", "0", "0.714", "1.0", "0.833"],
            ["salt_city_max_load", "10", "10", "10", "0", "0", "1.0", "1.0", "1.0"],
        ],
    )

    # ===== 第 10 章 =====
    add_heading(doc, "第 10 章　已知限制與風險", 1)
    add_heading(doc, "10.1 技術限制", 2)
    add_bullets(
        doc,
        [
            "Heuristic 抽取對複雜中文代詞、長句仍可能誤抽",
            "無獨立時間線引擎，僅有部分狀態先後規則",
            "規則覆蓋僅 7 條，無法涵蓋所有敘事矛盾類型",
            "HuggingFace provider 為 stub，尚未實作",
            "部分模組測試為 smoke level，非完整行為覆蓋",
        ],
    )
    add_heading(doc, "10.2 展示風險與備案", 2)
    add_table(
        doc,
        ["風險", "備案"],
        [
            ["Streamlit 啟動失敗", "先跑 pytest；展示已匯出 JSON"],
            ["Ollama 未安裝", "維持 heuristic 預設"],
            ["PyVis 渲染失敗", "展示 JSON stats + fallback"],
            ["某文本誤報", "改用 sample_clean_story.txt 對照"],
        ],
    )

    # ===== 第 11 章 =====
    add_heading(doc, "第 11 章　未來工作", 1)
    add_numbered(
        doc,
        [
            "FastAPI 後端：分離 UI 與分析 API，支援批次處理",
            "HuggingFace / llama.cpp：無 Ollama 時的本地推理替代",
            "獨立時間線引擎：章節序、事件先後、跳躍時間處理",
            "抽取品質：更完整的中文 NER、代詞消解、實體合併",
            "評估擴充：更多 ground truth、正式 benchmark 指標",
            "自動修稿：需另立作者確認與安全機制",
        ],
    )

    # ===== 第 12 章 =====
    add_heading(doc, "第 12 章　結論", 1)
    add_para(
        doc,
        "輕量智敘期末 MVP 已完成一套可本地執行、規則可解釋、結果可測試的中文敘事一致性檢查流程。"
        "系統以 Heuristic 抽取與 7 條 deterministic 規則為核心，Ollama 僅作可選輔助；"
        "在主要展示文本（霧城之鐘）上達成 Precision / Recall / F1 皆為 1.0，並通過 30 項 pytest 驗收。",
        indent=True,
    )
    add_para(
        doc,
        "本專題的定位是輔助作者發現敘事問題，而非取代創作判斷。"
        "誠實標示的限制與未實作項目，有助於後續迭代時聚焦於抽取品質、時間線推理與更廣泛的文本泛化。",
        indent=True,
    )

    # ===== 附錄 =====
    doc.add_page_break()
    add_heading(doc, "附錄 A　專案目錄結構", 1)
    add_code_block(
        doc,
        [
            "light-narrative/",
            "  app.py                      # Streamlit 入口",
            "  requirements.txt",
            "  samples/                    # 展示用範例文本",
            "  data/                       # SQLite 與 JSON 匯出",
            "  src/                        # 核心模組",
            "  tests/                      # 30 個 pytest",
            "  evaluation/                 # ground truth 評估模組",
        ],
    )

    add_heading(doc, "附錄 B　衝突類型中文對照", 1)
    add_table(
        doc,
        ["conflict_type", "UI 顯示"],
        [
            ["world_rule_violation", "世界規則違反"],
            ["character_state_conflict", "角色狀態衝突"],
            ["unique_item_conflict", "唯一物品衝突"],
            ["item_location_conflict", "物品位置衝突"],
            ["character_consistency_drift", "人設一致性漂移"],
            ["world_setting_conflict", "世界設定衝突"],
        ],
    )

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(
        "— 本報告依據 repo 實際程式碼、FINAL_STATUS.md、LIMITATIONS.md 與 evaluation/results/ 撰寫 —"
    )
    set_run_font(fr, 10, color=RGBColor(0x66, 0x66, 0x66))

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"已產生：{OUTPUT}")


if __name__ == "__main__":
    main()
