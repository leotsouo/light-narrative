"""讀取上傳的故事檔案。"""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument

from src.schemas import Document


def read_text_file(path: Path | str) -> str:
    path = Path(path)
    for encoding in ("utf-8", "utf-8-sig", "gbk", "big5"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def read_docx_bytes(data: bytes) -> str:
    doc = DocxDocument(BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def read_upload(filename: str, raw_bytes: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".docx"):
        return read_docx_bytes(raw_bytes)
    if lower.endswith((".txt", ".md", ".markdown")):
        for encoding in ("utf-8", "utf-8-sig", "gbk"):
            try:
                return raw_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        return raw_bytes.decode("utf-8", errors="replace")
    raise ValueError(f"不支援的檔案格式: {filename}")


def build_document(project_id: str, filename: str, text: str) -> Document:
    return Document(project_id=project_id, filename=filename, raw_text=text)
