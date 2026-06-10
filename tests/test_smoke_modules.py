from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument

from src.agents.report_agent import run_extraction
from src.chunker import chunk_document
from src.conflict_rules import detect_all_conflicts
from src.graph_builder import build_knowledge_graph, graph_stats
from src.ingest import read_docx_bytes, read_text_file, read_upload
from src.report_writer import build_report


def test_graph_builder_smoke() -> None:
    text = "## 第一章\n陳安停止呼吸。\n\n陳安從她身後說：「你好。」\n"
    chunks = chunk_document(text, "d1", "p1", strategy="chapter")
    extraction = run_extraction(chunks, use_llm=False)
    graph = build_knowledge_graph(extraction)
    stats = graph_stats(graph)
    assert stats["nodes"] >= 1


def test_report_writer_smoke() -> None:
    text = """## 第一章：死亡
羅恩停止呼吸。艾琳確認他已經死亡。

## 第二章：又出現
羅恩從她身後說：「不要害怕，孩子。」
"""
    chunks = chunk_document(text, "d1", "p1", strategy="chapter")
    extraction = run_extraction(chunks, use_llm=False)
    conflicts = detect_all_conflicts(extraction)
    report = build_report("p1", extraction, conflicts, chunk_count=len(chunks))
    assert report.chunk_count == len(chunks)
    assert report.entity_counts["characters"] >= 1
    assert len(report.conflicts) >= 1


def test_ingest_read_txt_md() -> None:
    sample = Path(__file__).resolve().parent.parent / "samples" / "sample_clean_story.txt"
    text = read_text_file(sample)
    assert "方遙" in text
    uploaded = read_upload("sample.md", sample.read_bytes())
    assert uploaded.replace("\r\n", "\n") == text.replace("\r\n", "\n")


def test_ingest_read_docx() -> None:
    doc = DocxDocument()
    doc.add_paragraph("第一章")
    doc.add_paragraph("方遙來到霧港。")
    buf = BytesIO()
    doc.save(buf)
    text = read_docx_bytes(buf.getvalue())
    assert "方遙" in text
